import requests, json
from django.shortcuts import render
from django.http import HttpResponse

import pandas as pd
import xlsxwriter
import io
import us
from pandas import json_normalize
from docx import Document

from ..forms import SearchForm
from .ct_view import *
from .lens_s_view import *
from .lens_p_view import *
from .nih_view import *


def authors_request(request):
    '''Queries information from clinicaltrials.gov API and lens.org API.

    arguments from form:
    
    used in both CT and Lens:
        author, institution, keyword
    used only in CT:
        sponsor
    
    '''
    if request.method == 'POST': # user has submitted data
        form = SearchForm(request.POST)
        if form.is_valid():
            fields = ['author', 'institution', 'sponsor', 'city', 'state', 'keyword', 'lens_id']
            entries = {}
            for field in fields:
                entry = form.cleaned_data[field]
                if entry:
                    entries[field] = entry

            ct_df = get_ct_df(entries)
            lens_s_df = get_lens_s_df(entries)
            lens_p_df = get_lens_p_df(entries)
            nih_df = get_nih_df(entries)

            author_doc = make_author_doc(ct_df, lens_s_df, lens_p_df, nih_df)

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=authors.docx'
            author_doc.save(response)

            return response
            
    else: # show empty form
        form = SearchForm()

    return render(request, 'search2.html', {'form': form})

def make_author_doc(ct_df, lens_s_df, lens_p_df, nih_df):
    '''Sort and filter results for common authors.

    arguments:
    df -- populated dataframe

    '''
    translator = {
        'Clinical Trials': [ct_df, 'Authors', 0, -1, "Title"],
        'Lens Scholar': [lens_s_df, 'AuthorNames', 0, -1, "Title"],
        'Lens Patent': [lens_p_df, 'inventors', 1, 0, "title"],
        'Federal NIH': [nih_df, 'piNames', 0, -1, "title"]
        }

    author_dict = populate_dict(ct_df, lens_s_df, lens_p_df, nih_df, translator)
    authors_sort_filt = sort_and_filter(author_dict)
    authors_document = make_doc(author_dict, authors_sort_filt)

    return authors_document

def populate_dict(ct_df, lens_s_df, lens_p_df, nih_df, translator):
    '''Creates dictionary mapping authors to work across different databases

    arguments:
    *_df: dataframe with data from the given database
    translator: info specific to each to database, such as author name arrangement, etc.
    
    '''
    author_dict = {}
    for db in translator:
        df, col_name, first, last, include = translator[db]
        for index, row in df.iterrows():
            if pd.isnull(df.loc[index, col_name]):
                continue
            entry = row[col_name]
            authors = entry.split(',')
            for author in authors:
                names = [n.strip() for n in author.split()]
                try:
                    name = (f'{names[first]} {names[last]}').title()
                except IndexError:
                    continue
                info = row[title].title()
                if name in author_dict:
                    if db in author_dict[name]:
                        author_dict[name][db].add(info)
                    else:
                        author_dict[name][db] = {info}
                else:
                    author_dict[name] = {db: {info}}

    return author_dict

def sort_and_filter(author_dict):
    authors_sort_filt = []
    for name in author_dict:
        categories = len(list(author_dict[name].keys()))
        total = sum([len(list(author_dict[name][key])) for key in author_dict[name]])
        if total > 1:
            authors_sort_filt.append((categories, total, name))
    authors_sort_filt.sort(reverse=True)

    return authors_sort_filt

def make_doc(author_dict, authors_sort_filt):
    document = Document()
    for _, _, name in authors_sort_filt:
        document.add_heading(name, level=1)
        data = author_dict[name]
        for database in data:
            document.add_heading(database, level=2)
            for title in data[database]:
                document.add_paragraph(title)
        document.add_page_break()
    
    return document